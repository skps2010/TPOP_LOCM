import subprocess
from os.path import exists
import docx
import json
from docx.text.font import Font
from docx.shared import Cm, Pt


def get_header_style(doc) -> Font:
    return doc.sections[0].header.paragraphs[0].runs[0].font


def get_body_style(doc) -> Font:
    return doc.paragraphs[0].runs[0].font


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def write_header(doc, lines: list, style: Font):
    for pr in doc.sections[0].header.paragraphs:
        delete_paragraph(pr)

    for l in lines:
        doc.sections[0].header.add_paragraph(l)

    for paragraph in doc.sections[0].header.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = style.color.rgb
            run.font.name = style.name
            run.font.size = style.size


def clear_body(doc):
    for pr in doc.paragraphs:
        delete_paragraph(pr)


def write_body(doc, lines: list, style: Font):
    for l in lines:
        doc.add_paragraph(l)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = style.color.rgb
            run.font.name = style.name
            run.font.size = style.size


def get_header_string():
    res = []
    authors: dict = json.loads(open("authors.json", "r").read())
    for n in authors["member"]:
        res.append(
            f"TPP2020-HW{authors['hw']}-{n['id']}{n['name']}")
    return res


def scan_file(fn: str) -> list:
    lines = [
        "",
        "/* -------------------------------------------------------------",
        f"// {fn}",
        "// ----------------------------------------------------------- */"
    ]
    with open(fn, "r") as fp:
        lines += fp.read().splitlines()
    return lines


if __name__ == "__main__":
    branch = subprocess.check_output(
        ['git', 'rev-parse', '--abbrev-ref',
         'HEAD']).decode().strip()
    target_code_files = subprocess.check_output(
        ['git', 'ls-tree', '-r', branch,
         '--name-only']).decode().splitlines()
    doc = docx.Document('tcc_template.docx')
    hs = get_header_style(doc)
    hs.size = Pt(10)

    bs = get_body_style(doc)
    write_header(doc, get_header_string(), hs)

    clear_body(doc)
    for file in target_code_files:
        print(f"processing file {file}")
        write_body(doc, scan_file(file), bs)

    doc.save("test.docx")
